import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import Pt
from io import BytesIO


def protokol_generator(bk_data, project):
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    TEMPLATE_FILE = os.path.join(BASE_DIR, 'protokoly', 'Wzor_pliku.docx')
    document = Document(TEMPLATE_FILE)
    # styles
    heading1 = document.styles['Heading 1']
    normal = document.styles['Normal']
    # print all avaiable styles
    # for style in document.styles:
    #     print("style.name == %s" % style.name)
    now = datetime.now()
    date_string = now.strftime("%d.%m.%Y")

    paragraph0_date_text = f"Pieńków, {date_string}"
    paragraph0_date = document.add_paragraph(
        paragraph0_date_text,
        style=normal,
    )
    # align to right
    paragraph0_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph0_date_run1 = paragraph0_date.add_run()
    paragraph0_date_run1.add_break()

    # add Title of document
    if "Numer ogłoszenia" in bk_data:
        paragraph0_title_text = f"Protokół z wyboru ofert do zapytania ofertowego nr {bk_data['Numer ogłoszenia']}"
    else:
        paragraph0_title_text = 'Protokół z wyboru ofert do zapytania ofertowego nr XX/XX/XXXX/2020'
    # create title
    paragraph0_title = document.add_paragraph(style=normal)
    # align center
    paragraph0_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # add run
    paragraph0_title_run0 = paragraph0_title.add_run()
    # make run bold
    paragraph0_title_run0.bold = True
    # add text
    paragraph0_title_run0.add_text(paragraph0_title_text)
    # set font size
    paragraph0_title_run0_font = paragraph0_title_run0.font
    paragraph0_title_run0_font.size = Pt(14)
    # add break
    paragraph0_title_run0.add_break()
    paragraph0_title_run0.add_break()

    # align to right

    ### I. Zamawiający
    paragraph1_header_text = 'Zamawiający'
    paragraph1_header = document.add_paragraph(
        paragraph1_header_text,
        style=heading1,
    )

    # I. Zamawiający Nazwa
    paragraph1_text = ""
    paragraph1 = document.add_paragraph(style=normal)
    if 'Nazwa' in bk_data:
        paragraph1_run1 = paragraph1.add_run(bk_data['Nazwa'])

    else:
        paragraph1_run1 = paragraph1.add_run('Adamed Pharma S.A.').bold = True
    paragraph1_run1.bold = True
    paragraph1_run1.add_break()
    # I. Zamawiający Adres
    if 'Adres' in bk_data:
        paragraph1_text += bk_data['Adres'] + ', \n'
    if 'Numer telefonu' in bk_data:
        paragraph1_text += f'tel.: {bk_data["Numer telefonu"]}' + ', \n'
    if 'NIP' in bk_data:
        paragraph1_text += f'NIP: {bk_data["NIP"]}' + ', \n'

    if not paragraph1_text:
        paragraph1_text = """
        Pieńków, ul. M. Adamkiewicza 6A, 05-152 Czosnów
        tel.: +48 513-000-101
        www.adamed.com.pl
        KRS: 0000116926
        NIP: 731-17-51-025
        """
    paragraph1.add_run(paragraph1_text)

    ### II. Informacje o projekcie
    paragraph2_header_text = 'Informacje o projekcie'
    paragraph2_header = document.add_paragraph(
        paragraph2_header_text,
        style=heading1,
    )
    paragraph2 = document.add_paragraph(style=normal)
    paragraph2_run1 = paragraph2.add_run('Tytuł projektu')
    paragraph2_run1.bold = True
    paragraph2_run1.add_break()
    if 'Tytuł projektu' in bk_data:
        paragraph2_run2 = paragraph2.add_run(bk_data['Tytuł projektu'])
        paragraph2_run2.add_break()
    paragraph2_run3 = paragraph2.add_run('\nNumer projektu')
    paragraph2_run3.bold = True
    paragraph2_run3.add_break()
    if 'Numer projektu' in bk_data:
        paragraph2_run4 = paragraph2.add_run(bk_data['Numer projektu'])
        paragraph2_run4.add_break()

    ### III.
    paragraph3_header_text = "Osoby wykonujące czynności związane z przeprowadzeniem postępowania o udzielenie zamówienia"
    paragraph3_header = document.add_paragraph(
        paragraph3_header_text,
        style=heading1,
    )

    if project == "LabGears":
        paragraph3_text = """
        1.	Adrian Jasiński - Ekspert
        2.	Tomasz Wiewióra – Kierownik Działu Finansowania Inwestycji
        """
    else:
        paragraph3_text = """
        1.	
        2.	
        3.
        """
    paragraph3 = document.add_paragraph(paragraph3_text, style=normal)

    ### IV.	Upublicznienie zapytania
    paragraph4_header_text = "Upublicznienie zapytania"
    paragraph4_header = document.add_paragraph(
        paragraph4_header_text,
        style=heading1,
    )
    if "Numer ogłoszenia" in bk_data and "Data publikacji" in bk_data:
        paragraph4_text = f"""
        Zapytanie ofertowe nr {bk_data['Numer ogłoszenia']} zostało upublicznione w dniu {bk_data['Data publikacji']} r. na stronie:

        https://bazakonkurencyjnosci.funduszeeuropejskie.gov.pl/
        """
    else:
        paragraph4_text = """
        Zapytanie ofertowe nr X/XX/XX/20XX zostało upublicznione w dniu XX.XX.XXXX r. na stronie:

        https://bazakonkurencyjnosci.funduszeeuropejskie.gov.pl/
        """
    paragraph4 = document.add_paragraph(paragraph4_text, style=normal)
    if 'Termin składania ofert' in bk_data:
        paragraph4_run1 = paragraph4.add_run()
        paragraph4_run1.add_break()
        paragraph4_run2 = paragraph4.add_run(
            f'Oferty należy złożyć {bk_data["Termin składania ofert"]}'
        )

    ### V.	Miejsce i sposób składania ofert
    paragraph5_header_text = "Miejsce i sposób składania ofert"
    paragraph5_header = document.add_paragraph(
        paragraph5_header_text,
        style=heading1,
    )

    paragraph5_text = bk_data['Miejsce i sposób składania ofert']
    paragraph5 = document.add_paragraph(paragraph5_text, style=normal)

    ### VI.	Przedmiot zamówienia
    paragraph6_header_text = "Przedmiot zamówienia"
    paragraph6_header = document.add_paragraph(
        paragraph6_header_text,
        style=heading1,
    )

    paragraph6_text = bk_data['Przedmiot zamówienia']
    paragraph6 = document.add_paragraph(paragraph6_text, style=normal)

    ### VII.	Zestawienie ofert
    paragraph7_header_text = "Zestawienie ofert"
    paragraph7_header = document.add_paragraph(
        paragraph7_header_text,
        style=heading1,
    )

    paragraph7_text = "W odpowiedzi na zapytanie ofertowe wpłynęły następujące oferty:"
    paragraph7 = document.add_paragraph(paragraph7_text, style=normal)

    ### VIII.	Ocena spełnienia warunków udziału w postępowaniu
    paragraph8_header_text = "Ocena spełnienia warunków udziału w postępowaniu"
    paragraph8_header = document.add_paragraph(
        paragraph8_header_text,
        style=heading1,
    )

    paragraph8 = document.add_paragraph(style=normal)
    paragraph8 = document.add_paragraph(style=normal)
    if "Wiedza i doświadczenie" in bk_data:
        paragraph8_run1 = paragraph8.add_run('Wiedza i doświadczenie')
        paragraph8_run1.bold = True
        paragraph8_run1.add_break()
        paragraph8_run2 = paragraph8.add_run(bk_data['Wiedza i doświadczenie'])
        paragraph8_run2.add_break()
        paragraph8_run2.add_break()
    if "Potencjał techniczny" in bk_data:
        paragraph8_run3 = paragraph8.add_run("Potencjał techniczny")
        paragraph8_run3.bold = True
        paragraph8_run3.add_break()
        paragraph8_run4 = paragraph8.add_run(bk_data["Potencjał techniczny"])
        paragraph8_run4.add_break()
        paragraph8_run4.add_break()
    if "Osoby zdolne do wykonania zamówienia" in bk_data:
        paragraph8_run5 = paragraph8.add_run(
            "Osoby zdolne do wykonania zamówienia"
        )
        paragraph8_run5.bold = True
        paragraph8_run5.add_break()
        paragraph8_run6 = paragraph8.add_run(
            bk_data["Osoby zdolne do wykonania zamówienia"]
        )
        paragraph8_run6.add_break()
        paragraph8_run6.add_break()
    if "Sytuacja ekonomiczna i finansowa" in bk_data:
        paragraph8_run7 = paragraph8.add_run("Sytuacja ekonomiczna i finansowa")
        paragraph8_run7.bold = True
        paragraph8_run7.add_break()
        paragraph8_run8 = paragraph8.add_run(
            bk_data["Sytuacja ekonomiczna i finansowa"]
        )
        paragraph8_run8.add_break()
        paragraph8_run8.add_break()
    if "Dodatkowe warunki" in bk_data:
        paragraph8_run9 = paragraph8.add_run("Dodatkowe warunki")
        paragraph8_run9.bold = True
        paragraph8_run9.add_break()
        paragraph8_run10 = paragraph8.add_run(bk_data["Dodatkowe warunki"])
        paragraph8_run10.add_break()
        paragraph8_run10.add_break()

    paragraph8_run11 = paragraph8.add_run("Ocena spełnienia warunków:")
    paragraph8_run11.bold = True
    paragraph8_run11.add_break()
    paragraph8_run11.add_break()

    ### IX.	Ocena wymaganych oświadczeń i dokumentów
    paragraph9_header_text = "Ocena wymaganych oświadczeń i dokumentów"
    paragraph9_header = document.add_paragraph(
        paragraph9_header_text,
        style=heading1,
    )

    paragraph9 = document.add_paragraph(style=normal)
    paragraph9_run1 = paragraph9.add_run(
        "Wykaz wymaganych oświadczeń lub dokumentów:"
    )
    paragraph9_run1.bold = True
    paragraph9_run1.add_break()

    if 'Lista dokumentów/oświadczeń wymaganych od Wykonawcy' in bk_data:
        paragraph9_run2 = paragraph9.add_run(
            bk_data['Lista dokumentów/oświadczeń wymaganych od Wykonawcy']
        )

    paragraph9_run2 = paragraph9.add_run()
    paragraph9_run2.add_break()
    paragraph9_run2.add_break()
    paragraph9_run2.add_text("Ocena spełnienia warunków:")
    paragraph9_run2.bold = True
    paragraph9_run2.add_break()

    ### X.	Kryteria oceny
    paragraph10_header_text = "Kryteria oceny"
    paragraph10_header = document.add_paragraph(
        paragraph10_header_text,
        style=heading1,
    )

    paragraph10 = document.add_paragraph(style=normal)
    paragraph10_run1 = paragraph10.add_run(
        "Zamawiający dokonał oceny złożonych ofert w oparciu o poniższe \
            kryteria, wagi procentowe oraz według określonego wzoru:"
    )
    paragraph10_run1.add_break()

    if 'Kryteria oceny i opis sposobu przyznawania punktacji' in bk_data:
        paragraph10_run2 = paragraph10.add_run(
            bk_data['Kryteria oceny i opis sposobu przyznawania punktacji']
        )

    ### XI.	Wykaz ofert dopuszczonych do oceny w postępowaniu ofertowym na podstawie przyjętych kryteriów
    paragraph11_header_text = "Wykaz ofert dopuszczonych do oceny w postępowaniu ofertowym na podstawie przyjętych kryteriów"
    paragraph11_header = document.add_paragraph(
        paragraph11_header_text,
        style=heading1,
    )

    ### XII.	Ocena ofert na podstawie określonych w zapytaniu kryteriów oceny ofert
    paragraph12_header_text = "Ocena ofert na podstawie określonych w zapytaniu kryteriów oceny ofert"
    paragraph12_header = document.add_paragraph(
        paragraph12_header_text,
        style=heading1,
    )
    ### XIII.	Podsumowanie
    paragraph13_header_text = "Podsumowanie / Wybór wykonawcy"
    paragraph13_header = document.add_paragraph(
        paragraph13_header_text,
        style=heading1,
    )

    # Jeżeli są informacje o wybranym wykonwcy
    if 'Nazwa i adres, data wpłynięcia oferty oraz jej cena' in bk_data:
        paragraph13 = document.add_paragraph(style=normal)
        if "Informacja o wybranym wykonawcy" in bk_data:
            paragraph13_run0 = paragraph13.add_run(
                bk_data["Informacja o wybranym wykonawcy"]
            )
            paragraph13_run0.add_break()
            paragraph13_run0.add_break()

        paragraph13_run1 = paragraph13.add_run(
            'Nazwa i adres, data wpłynięcia oferty oraz jej cena'
        )
        paragraph13_run1.bold = True
        paragraph13_run1.add_break()

        paragraph13_run2 = paragraph13.add_run(
            bk_data['Nazwa i adres, data wpłynięcia oferty oraz jej cena']
        )

    ### XIV. Informacje o Braku Powiązań
    if project == "LabGears":
        add_oswiadczenie(
            document,
            bk_data,
            project,
            "Adrian Jasiński",
            "Ekpert",
        )
        add_oswiadczenie(
            document,
            bk_data,
            project,
            "Tomasz Wiewióra",
            "Kierownik Działu Finansowania Inwestycji",
        )
    else:
        add_oswiadczenie(document, bk_data, project, "   ", "   ")
        add_oswiadczenie(document, bk_data, project, "   ", "   ")
    # RETURN DOCX FILE
    # define filename
    filename = f"{bk_data['Numer ogłoszenia']}_{date_string.replace('.', '_')}.docx"
    # define in memoty bytes structure
    bio = BytesIO()
    # save document in memory
    document.save(bio)
    # rewind the stream
    bio.seek(0)
    # return filename and file bytes structure
    return (filename, bio.getvalue())


def add_oswiadczenie(document, bk_data, project, imie_nazwisko, stanowisko):
    # dodaj przerwe aby kolejne informacje były na nowej stronie
    document.add_page_break()
    normal = document.styles['Normal']
    list2 = document.styles['List Paragraph']
    now = datetime.now()
    date_string = now.strftime("%d.%m.%Y")

    paragraph0_date_text = f"Pieńków, {date_string}"
    paragraph0_date = document.add_paragraph(
        paragraph0_date_text,
        style=normal,
    )
    # align to right
    paragraph0_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph0_date_run1 = paragraph0_date.add_run()
    paragraph0_date_run1.add_break()

    # add page top
    paragraph14_top = document.add_paragraph(style=normal)
    # add run
    paragraph14_top_run0 = paragraph14_top.add_run()
    # add text
    paragraph14_top_run0.add_text(imie_nazwisko)
    # set font size
    paragraph14_top_run0.add_break()

    paragraph14_top_run1 = paragraph14_top.add_run()
    paragraph14_top_run1.add_text("Imię i nazwisko")
    paragraph14_top_run1_font = paragraph14_top_run1.font
    paragraph14_top_run1_font.size = Pt(8)
    paragraph14_top_run1.add_break()

    paragraph14_top_run2 = paragraph14_top.add_run()
    paragraph14_top_run2.add_text(stanowisko)
    paragraph14_top_run2.add_break()

    paragraph14_top_run3 = paragraph14_top.add_run()
    paragraph14_top_run3_font = paragraph14_top_run3.font
    paragraph14_top_run3_font.size = Pt(8)
    paragraph14_top_run3.add_text("Stanowisko")
    paragraph14_top_run3.add_break()

    paragraph14_header_text = "Oświadczenie o braku powiązań osobowych lub kapitałowych"
    paragraph14_header = document.add_paragraph(style=normal)
    paragraph14_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph14_header_run = paragraph14_header.add_run()
    paragraph14_header_run.add_text(paragraph14_header_text)
    paragraph14_header_run.bold = True
    paragraph14_header_run.add_break()

    paragraph14_body1 = document.add_paragraph(style=normal)
    paragraph14_body1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph14_body1_run0 = paragraph14_body1.add_run()
    paragraph14_body1_run0_text = f"""
    Ja niżej podpisana, oświadczam, że nie jestem powiązana osobowo lub kapitałowo z żadnym z oferentów, którzy złożyli swoje oferty w odpowiedzi na zapytanie ofertowe nr {bk_data['Numer ogłoszenia']}.
    """
    paragraph14_body1_run0.add_text(paragraph14_body1_run0_text)
    paragraph14_body1_run0.add_break()

    paragraph14_body1_run1 = paragraph14_body1.add_run()

    paragraph14_body1_run1_text = """
    Przez powiązania kapitałowe lub osobowe rozumie się wzajemne powiązania między Zamawiającym lub osobami upoważnionymi do zaciągania zobowiązań w imieniu Zamawiającego
     lub osobami wykonującymi w imieniu Zamawiającego czynności związane z przygotowaniem i przeprowadzeniem procedury wyboru wykonawcy a wykonawcą, polegające w szczególności na:\n
    """
    paragraph14_body1_run1.add_text(paragraph14_body1_run1_text)
    paragraph14_body1_run1.add_break()

    paragraph14_body2 = document.add_paragraph(style=list2)
    paragraph_format = paragraph14_body2.paragraph_format
    paragraph_format.left_indent
    paragraph_format.left_indent = Inches(0.7)

    paragraph14_body2_run0 = paragraph14_body2.add_run()
    paragraph14_body2_run0.add_text(
        "a)	  uczestniczeniu w spółce jako wspólnik spółki cywilnej lub spółki osobowej,"
    )
    paragraph14_body2_run0.add_break()
    paragraph14_body2_run1 = paragraph14_body2.add_run()
    paragraph14_body2_run1.add_text(
        "b)	  posiadaniu co najmniej 10 % udziałów lub akcji,"
    )
    paragraph14_body2_run1.add_break()
    paragraph14_body2_run2 = paragraph14_body2.add_run()
    paragraph14_body2_run2.add_text(
        "c)	  pełnieniu funkcji członka organu nadzorczego lub zarządzającego, prokurenta, pełnomocnika,"
    )
    paragraph14_body2_run2.add_break()
    paragraph14_body2_run3 = paragraph14_body2.add_run()
    paragraph14_body2_run3.add_text(
        """d)   pozostawaniu w związku małżeńskim, w stosunku pokrewieństwa lub powinowactwa w linii prostej, pokrewieństwa lub powinowactwa w linii bocznej do drugiego stopnia lub w stosunku przysposobienia, opieki lub kurateli."""
    )
    paragraph14_body2_run3.add_break()

    paragraph14_bottom = document.add_paragraph(style=normal)
    paragraph14_bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph14_bottom_run0 = paragraph14_bottom.add_run()
    paragraph14_bottom_run0.add_break()
    paragraph14_bottom_run0.add_break()
    paragraph14_bottom_run0.add_break()
    paragraph14_bottom_run0.add_break()
    paragraph14_bottom_run0.add_break()
    paragraph14_bottom_run0.add_break()

    paragraph14_bottom_run0.add_text("...............................")
    paragraph14_bottom_run0.add_break()

    paragraph14_bottom_run1 = paragraph14_bottom.add_run()
    paragraph14_bottom_run1.add_text("data i podpis")
    paragraph14_bottom_run1_font = paragraph14_top_run1.font
    paragraph14_bottom_run1_font.size = Pt(8)

    return document
